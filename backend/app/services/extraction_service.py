from __future__ import annotations

import io
from pathlib import Path

from app.core.logging import get_logger

logger = get_logger("digitlaw.extraction")


def extract_text(file_path: str, filename: str) -> str:
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(file_path)
    if suffix == ".docx":
        return _extract_docx(file_path)

    raise ValueError(f"Unsupported file type: {suffix}")


def _extract_pdf(file_path: str) -> str:
    import pypdf

    text_parts = []
    with open(file_path, "rb") as f:
        reader = pypdf.PdfReader(f)
        for page in reader.pages:
            text_parts.append(page.extract_text() or "")
    return "\n".join(text_parts)


def _extract_docx(file_path: str) -> str:
    import docx

    document = docx.Document(file_path)
    parts = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))

    return "\n".join(parts)
